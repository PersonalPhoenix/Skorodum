import io
import os

from rest_framework.generics import (
    DestroyAPIView, 
    UpdateAPIView,
    ListAPIView,
)
from rest_framework.views import (
    APIView,
)
from rest_framework.decorators import (
    api_view,
)
from rest_framework.response import (
    Response,
)
from rest_framework import (
    status,
)
from django.http import (
    HttpResponse,
)
from django.db.models import (
    Q,
)
from django.forms.models import (
    model_to_dict,
)
from django.http import (
    HttpResponse,
)
from django.conf import (
    settings,
)
from django.http import (
    StreamingHttpResponse,
)
from wsgiref.util import (
    FileWrapper,
)
from docx import (
    Document,
)

from .models import (
    Game,
    Round,
    Question,
    Category,
)
from .serializers import (
    CategorySerializer,
    QuestionSerializer,
    FileSerializer,
    RoundSerializer,
)
from .helpers.game_helpers import (
    create_game,
    get_names_all_games,
    get_one_game_with_rounds,
    update_game,
    get_game_for_json,
    create_custom_style,
)
from .helpers.round_helpers import (
    create_round,
    update_round,
)
from .helpers.question_helpers import (
    create_question,
    get_one_question,
    update_question,
)


class GameCreateAPI(APIView):

    def post(self, request, format=None):
        obj, create = create_game(request.data)

        if create:
            return Response(obj, status=status.HTTP_201_CREATED)

        return Response(obj, status=status.HTTP_400_BAD_REQUEST)


class GameGetAPI(APIView):

    def get(self, request, *args, **kwargs):
        result = get_one_game_with_rounds(request, *args, **kwargs)

        return Response(result, status=status.HTTP_200_OK)


class GameGetNamesAPI(APIView):

    def get(self, request, *args, **kwargs):
        return Response(get_names_all_games(), status=status.HTTP_200_OK)


@api_view(['GET'])
def download_file(request):
    file_name = request.GET.get('file_name')
    if file_name:
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        if os.path.isfile(file_path):
            wrapper = FileWrapper(open(file_path, 'rb'))
            content_type = 'application/force-download'
            response = HttpResponse(wrapper, content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{file_name}"'
            return response
        else:
            return HttpResponse(status=404)
    else:
        return HttpResponse(status=400)


@api_view(['POST'])
def upload_file(request):
    file_serializer = FileSerializer(data=request.data)
    if file_serializer.is_valid():
        file_serializer.save()

        return Response(file_serializer.data, status=status.HTTP_201_CREATED)

    return Response(file_serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class GameUpdateAPI(UpdateAPIView):

    def put(self, request, *args, **kwargs):

        Game.objects.filter(id=kwargs['pk']).update(**request.data)

        return HttpResponse('Ok', status=200)


class GameDeleteAPI(DestroyAPIView):

    serializer_class = QuestionSerializer
    queryset = Game.objects.all()


class GameUploadJsonAPI(APIView):

    def get(self, request, *args, **kwargs):
        game, game_id = get_game_for_json(request, *args, **kwargs)
        response = Response(game, content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename="Игра №{game_id}.json"'

        return response


class GameUploadWordAPI(APIView):

    def get(self, request, *args, **kwargs):
        data = get_one_game_with_rounds(request, *args, **kwargs)
        document = Document()

        custom_main_header, \
        custom_sub_header, \
        custom_info = create_custom_style(document)

        document.add_heading(f'{data["name"]}').style = custom_main_header
        document.add_heading(f'Тема: {data["theme"]}').style = custom_sub_header

        document.add_paragraph(f'Клиент: {data["client"]}').style = custom_info
        document.add_paragraph(f'Дата создания: {data["date"]}').style = custom_info
        document.add_paragraph(f'Удалить ответ: {data["remove_answer"]}').style = custom_info
        document.add_paragraph(f'Один за всех: {data["one_for_all"]}').style = custom_info
        document.add_paragraph(f'Вопрос со ставкой: {data["question_bet"]}').style = custom_info
        document.add_paragraph(f'All-in: {data["all_in"]}').style = custom_info
        document.add_paragraph(f'Командная ставка: {data["team_bet"]}').style = custom_info
        document.add_paragraph(f'Пропустить emails: {data["skip_emails"]}').style = custom_info

        for round_number, _round,  in enumerate(data['rounds'], start=1):
            document.add_heading(f'Раунд № {round_number} - {_round["name"]}').style = custom_sub_header
            document.add_heading(f'Тип раунда: {_round["round_type"]}').style = custom_sub_header

            document.add_paragraph(f'Тестовый раунд: {"Да" if _round["is_test"] else "Нет"}').style = custom_info
            document.add_paragraph(f'Отображение имен на экране: {"Да" if _round["display_name"] else "Нет"} ').style = custom_info
            document.add_paragraph(f'Время ответа на вопрос: {_round["time_to_answer"]} сек').style = custom_info
            document.add_paragraph(f'Использование специальных тактик: {"Да" if _round["use_special_tactics"] else "Нет"}').style = custom_info

            for question_number, question in enumerate(_round['questions'], start=1):
                document.add_heading(f'Вопрос № {question_number} Раунда № {round_number}').style = custom_sub_header
                document.add_heading(f'Тип вопроса: {question["question_type"]}').style = custom_sub_header

                document.add_paragraph(f'Категории вопроса: {question["category_names"]}').style = custom_info
                document.add_paragraph(f'Время ответа на вопрос: {question["question_text"]}').style = custom_info
                document.add_paragraph(f'Показывать изображение: {"Да" if question["show_image"] else "Нет"}').style = custom_info
                document.add_paragraph(f'Изображение до вопроса: {question["image_before"]}').style = custom_info
                document.add_paragraph(f'Изображение после вопроса: {question["image_after"]}').style = custom_info
                document.add_paragraph(f'Видео до вопроса: {question["video_before"]}').style = custom_info
                document.add_paragraph(f'Видео после вопроса: {question["video_after"]}').style = custom_info
                document.add_paragraph(f'Ответы на вопрос: {question["answers"]}').style = custom_info
                document.add_paragraph(f'Отображать на экране: {"Да" if question["player_displayed"] else "Нет"}').style = custom_info

        buffer = io.BytesIO()
        document.save(buffer) 
        buffer.seek(0)  

        response = StreamingHttpResponse(
            streaming_content=buffer,  
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingm'
        )

        response['Content-Disposition'] = 'attachment;filename=Test.docx'
        response["Content-Encoding"] = 'UTF-8'

        return response


class RoundCreateAPI(APIView):

    def post(self, request, format=None):
        obj, create = create_round(request.data)

        if create:
            return Response(obj, status=status.HTTP_201_CREATED)

        return Response(obj, status=status.HTTP_400_BAD_REQUEST)


class RoundUpdateAPI(UpdateAPIView):

    serializer_class = RoundSerializer
    queryset = Round.objects.all()


class RoundDeleteAPI(DestroyAPIView):

    serializer_class = QuestionSerializer
    queryset = Round.objects.all()


class QuestionCreateAPI(APIView):
    
    def post(self, request, format=None):
        obj, create = create_question(request.data)

        if create:
            return Response(obj, status=status.HTTP_201_CREATED)

        return Response(obj, status=status.HTTP_400_BAD_REQUEST)


class QuestionGetOneAPI(APIView):

    def get(self, request, *args, **kwargs):
        return Response(get_one_question(request.data, *args, **kwargs), status=status.HTTP_200_OK)


class QuestionGetAllAPI(ListAPIView):

    def get(self, request, *args, **kwargs):
        return Response(Question.objects.filter().values(), status=status.HTTP_200_OK)


class QuestionUpdateAPI(UpdateAPIView):

    def put(self, request, *args, **kwargs):
        instance = Question.objects.get(id=kwargs['pk'])
        serializer = QuestionSerializer(instance, data=request.data)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, *args, **kwargs):
        instance = Question.objects.get(kwargs['pk'])
        serializer = QuestionSerializer(instance, data=request.data, partial=True)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class QuestionDeleteAPI(APIView):

    def delete(self, request, pk):
        try:
            question = Question.objects.get(pk=pk)
        except Question.DoesNotExist:
            return Response({"detail": "Вопрос не найден."}, status=status.HTTP_404_NOT_FOUND)

        question.round_id.clear()
        question.category.clear()
        question.delete()

        return Response(status=status.HTTP_204_NO_CONTENT)


class QuestionSearchAPI(ListAPIView):
    serializer_class = QuestionSerializer

    def get_queryset(self):
        query = self.request.query_params.get('query', '')

        return Question.objects.filter(
            Q(question_text__icontains=query),
        )

class CategoryCreateAPI(APIView):
    
    def post(self, request, format=None):
        serializer = CategorySerializer(data=request.data)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CategoryGetAllAPI(APIView):
    
    def get(self, request, *args, **kwargs):
        result = Category.objects.all().order_by('-id').values()

        return Response(result, status=status.HTTP_200_OK)


class CategoryGetAPI(APIView):

    def get(self, request, *args, **kwargs):
        result = Category.objects.values().get(id=kwargs['pk'])

        return Response(result, status=status.HTTP_200_OK)


class CategoryGetKeysAPI(APIView):

    def get(self, request, *args, **kwargs):
        ids = kwargs['ids'].split(',')
        result = Category.objects.filter(id__in=ids).values()

        return Response(result, status=status.HTTP_200_OK)


class CategorySearchAPI(ListAPIView):
    serializer_class = CategorySerializer

    def get_queryset(self):
        query = self.request.query_params.get('query', '')

        return Category.objects.filter(
            Q(name__icontains=query),
        )


class CategoryUpdateAPI(UpdateAPIView):

    serializer_class = CategorySerializer
    queryset = Category.objects.all()


class CategoryDeleteAPI(DestroyAPIView):

    serializer_class = CategorySerializer
    queryset = Category.objects.all()
