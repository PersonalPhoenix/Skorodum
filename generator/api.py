import io
import os

from rest_framework.generics import (
    DestroyAPIView, 
    UpdateAPIView,
    ListAPIView,
)
import json
import zipfile

from rest_framework.exceptions import (
    NotFound,
)
from rest_framework.views import (
    APIView,
)
from rest_framework.decorators import (
    api_view,
)
from rest_framework.response import (
    Response,
)
from rest_framework import (
    status,
)
from django.http import (
    HttpResponse,
)
from django.db.models import (
    Q,
)
from wsgiref.util import (
    FileWrapper,
)
from django.http import (
    HttpResponse,
)
from django.conf import (
    settings,
)
from django.http import (
    StreamingHttpResponse,
)
from docx import (
    Document,
)

from .models import (
    Game,
    Round,
    Question,
    Category,
)
from .serializer import (
    CategorySerializer,
    QuestionSerializer,
    FileSerializer,
    RoundSerializer,
    GameSerializer,
)
from .helpers.game_helpers import (
    create_game,
    get_names_all_games,
    get_one_game_with_rounds,
    get_game_for_json,
    create_custom_style,
    get_selected_games,
    get_media,
)
from .helpers.round_helpers import (
    create_round,
    get_one_round,
)
from .helpers.question_helpers import (
    create_question,
    get_one_question,
)


class GameCreateAPI(APIView):

    def post(self, request, format=None):
        obj, create = create_game(request.data)

        if create:
            return Response(obj, status=status.HTTP_201_CREATED)

        return Response(obj, status=status.HTTP_400_BAD_REQUEST)


class GameGetAPI(APIView):

    def get(self, request, *args, **kwargs):
        try:
            result = get_one_game_with_rounds(request, *args, **kwargs)
            return Response(result, status=status.HTTP_200_OK)

        except Game.DoesNotExist:
            raise NotFound(detail=f'Игра с id: {kwargs['pk']} не существует.')


class GameGetNamesAPI(APIView):

    def get(self, request, *args, **kwargs):
        return Response(get_names_all_games(), status=status.HTTP_200_OK)


@api_view(['GET'])
def download_file(request):
    file_name = request.GET.get('file_name')
    if file_name:
        file_path = os.path.join(settings.MEDIA_ROOT, file_name)
        if os.path.isfile(file_path):
            wrapper = FileWrapper(open(file_path, 'rb'))
            content_type = 'application/force-download'
            response = HttpResponse(wrapper, content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{file_name}"'
            return response
        else:
            return HttpResponse(status=404)
    else:
        return HttpResponse(status=400)


@api_view(['POST'])
def upload_file(request):
    file_serializer = FileSerializer(data=request.data)
    if file_serializer.is_valid():
        file = file_serializer.save()

        return Response(file_serializer.data, status=status.HTTP_201_CREATED)

    return Response(file_serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class GameUpdateAPI(UpdateAPIView):

    serializer_class = GameSerializer
    queryset = Game.objects.all()


class GameDeleteAPI(DestroyAPIView):

    def delete(self, request, *args, **kwargs):
        try:
            Game.objects.get(id=kwargs['pk']).delete()
            return HttpResponse(f'Игра с id: {kwargs['pk']} была удалена', status=200)
        except Game.DoesNotExist:
            raise NotFound(detail=f'Игра с id: {kwargs['pk']} не существует.')


class GameUploadJsonAPI(APIView):

    def get(self, request, *args, **kwargs):
        game, game_id = get_game_for_json(request, *args, **kwargs)
        response = Response(game, content_type='application/json')
        response['Content-Disposition'] = f'attachment; filename="Игра №{game_id}.json"'

        return response


class GameUploadWordAPI(APIView):

    def get(self, request, *args, **kwargs):
        data = get_one_game_with_rounds(request, *args, **kwargs)
        document = Document()

        custom_main_header, \
        custom_sub_header, \
        custom_info = create_custom_style(document)

        document.add_heading(f'{data["name"]}').style = custom_main_header
        document.add_heading(f'Тема: {data["theme"]}').style = custom_sub_header

        document.add_paragraph(f'Клиент: {data["client"]}').style = custom_info
        document.add_paragraph(f'Дата создания: {data["date"]}').style = custom_info
        document.add_paragraph(f'Удалить ответ: {data["remove_answer"]}').style = custom_info
        document.add_paragraph(f'Один за всех: {data["one_for_all"]}').style = custom_info
        document.add_paragraph(f'Вопрос со ставкой: {data["question_bet"]}').style = custom_info
        document.add_paragraph(f'All-in: {data["all_in"]}').style = custom_info
        document.add_paragraph(f'Командная ставка: {data["team_bet"]}').style = custom_info
        document.add_paragraph(f'Пропустить emails: {data["skip_emails"]}').style = custom_info

        for round_number, _round,  in enumerate(data['rounds'], start=1):
            document.add_heading(f'Раунд № {round_number} - {_round["name"]}').style = custom_sub_header
            document.add_heading(f'Тип раунда: {_round["round_type"]}').style = custom_sub_header

            document.add_paragraph(f'Тестовый раунд: {"Да" if _round["is_test"] else "Нет"}').style = custom_info
            document.add_paragraph(f'Отображение имен на экране: {"Да" if _round["display_name"] else "Нет"} ').style = custom_info
            document.add_paragraph(f'Время ответа на вопрос: {_round["time_to_answer"]} сек').style = custom_info
            document.add_paragraph(f'Использование специальных тактик: {"Да" if _round["use_special_tactics"] else "Нет"}').style = custom_info

            for question_number, question in enumerate(_round['questions'], start=1):
                document.add_heading(f'Вопрос № {question_number} Раунда № {round_number}').style = custom_sub_header
                document.add_heading(f'Тип вопроса: {question["question_type"]}').style = custom_sub_header

                document.add_paragraph(f'Категории вопроса: {question["category_names"]}').style = custom_info
                document.add_paragraph(f'Время ответа на вопрос: {question["question_text"]}').style = custom_info
                document.add_paragraph(f'Показывать изображение: {"Да" if question["show_image"] else "Нет"}').style = custom_info
                document.add_paragraph(f'Изображение до вопроса: {question["image_before"]}').style = custom_info
                document.add_paragraph(f'Изображение после вопроса: {question["image_after"]}').style = custom_info
                document.add_paragraph(f'Видео до вопроса: {question["video_before"]}').style = custom_info
                document.add_paragraph(f'Видео после вопроса: {question["video_after"]}').style = custom_info
                document.add_paragraph(f'Ответы на вопрос: {question["answers"]}').style = custom_info
                document.add_paragraph(f'Отображать на экране: {"Да" if question["player_displayed"] else "Нет"}').style = custom_info

        buffer = io.BytesIO()
        document.save(buffer) 
        buffer.seek(0)  

        response = StreamingHttpResponse(
            streaming_content=buffer,  
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingm'
        )

        response['Content-Disposition'] = 'attachment;filename=Test.docx'
        response["Content-Encoding"] = 'UTF-8'

        return response


class GameGetSelectedAPI(APIView):
    def get(self, request, *args, **kwargs):
        return Response(get_selected_games(kwargs.get('ids', '').split(',')), status=status.HTTP_200_OK)


class GameUploadZipAPI(APIView):
    def get(self, request, *args, **kwargs):
        ids = kwargs.get('ids', '').split(',')
        games = get_selected_games(ids)

        response = HttpResponse(content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="{len(ids)} games.zip"'

        with zipfile.ZipFile(response, 'w') as zipf:
            for game_id, game_data in games.items():
                json_data = json.dumps(game_data, indent=4, ensure_ascii=False).encode('utf-8')
                zipf.writestr(f'game{game_id}.json', json_data)

                medias = get_media(int(game_id))
                for media in medias:
                    for _, value in media.items():
                        md_path = os.path.join(settings.MEDIA_ROOT, value)
                        zipf.write(md_path, os.path.basename(md_path))

        return response


# ------------------------------------
# ROUND
# ------------------------------------

class RoundCreateAPI(APIView):

    def post(self, request, format=None):
        obj, create = create_round(request.data)

        if create:
            return Response(obj, status=status.HTTP_201_CREATED)

        return Response(obj, status=status.HTTP_400_BAD_REQUEST)


class RoundGetAPI(APIView):

    def get(self, request, *args, **kwargs):

        try:
            result = get_one_round(request, *args, **kwargs)
            return Response(result, status=status.HTTP_200_OK)

        except Round.DoesNotExist:
            raise NotFound(detail=f'Раунд с id: {kwargs['pk']} не существует')


class RoundUpdateAPI(UpdateAPIView):

    serializer_class = RoundSerializer
    queryset = Round.objects.all()


class RoundDeleteAPI(DestroyAPIView):

    def delete(self, request, *args, **kwargs):
        try:
            Round.objects.get(id=kwargs['pk']).delete()
            return Response({'detail':f'Раунд с id: {kwargs['pk']} был удален'}, status=200)
        except Round.DoesNotExist:
            raise NotFound(detail=f'Раунд с id: {kwargs['pk']} не существует')


# ------------------------------------
# QUESTION
# ------------------------------------


@api_view(['POST'])
def create_question(request):
    serializer = QuestionSerializer(data=request.data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class QuestionGetOneAPI(APIView):

    def get(self, request, *args, **kwargs):
        return Response(get_one_question(request.data, *args, **kwargs), status=status.HTTP_200_OK)


class QuestionGetAllAPI(ListAPIView):

    def get(self, request, *args, **kwargs):
        return Response(Question.objects.filter().values(), status=status.HTTP_200_OK)


class QuestionUpdateAPI(UpdateAPIView):

    def put(self, request, *args, **kwargs):
        instance = Question.objects.get(id=kwargs['pk'])
        serializer = QuestionSerializer(instance, data=request.data)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def patch(self, request, *args, **kwargs):
        instance = Question.objects.get(kwargs['pk'])
        serializer = QuestionSerializer(instance, data=request.data, partial=True)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class QuestionDeleteAPI(APIView):

    def delete(self, request, pk):
        try:
            question = Question.objects.get(pk=pk)
        except Question.DoesNotExist:
            return Response({'detail': f'Вопрос c id: {pk} не существует.'}, status=status.HTTP_404_NOT_FOUND)

        question.round_id.clear()
        question.category.clear()
        question.delete()

        return Response(status=status.HTTP_204_NO_CONTENT)


class QuestionSearchAPI(ListAPIView):
    serializer_class = QuestionSerializer

    def get_queryset(self):
        query = self.request.query_params.get('query', '')

        return Question.objects.filter(
            Q(question_text__icontains=query),
        )


class QuestionWordAPI(APIView):

    def get(self, request, *args, **kwargs):
        ids = kwargs.get('ids', '')

        mapping = [
            'question_type',
            'category__name',
            'question_text',
            'show_image',
            'image_before',
            'image_after',
            'video_before',
            'video_after',
            'player_displayed',
            'time_to_answer',
            'answers',
            'correct_answer',
            'open_question',
            'close_question',
            'media_question',
        ]
        if ids:
            data = Question.objects.filter(id__in=ids.split(',')).select_related('category').values(
                *mapping,
            )
        else:
            data = Question.objects.all().select_related('category').values(
                *mapping,
            )

        document = Document()

        for question in data:
            document.add_heading(f'Тип вопроса: {question["question_type"]}')
            document.add_heading(f'Категория: {question["category__name"]}')

            document.add_paragraph(f'Вопрос: {question["question_text"]}')
            document.add_paragraph(str(question['show_image']))
            document.add_paragraph(question['image_before'])
            document.add_paragraph(question['image_after'])
            document.add_paragraph(question['video_before'])
            document.add_paragraph(question['video_after'])
            document.add_paragraph(str(question['player_displayed']))
            document.add_paragraph(str(question['time_to_answer']))
            document.add_paragraph(question['answers'])
            document.add_paragraph(question['correct_answer'])
            document.add_paragraph(str(question['open_question']))
            document.add_paragraph(str(question['close_question']))
            document.add_paragraph(str(question['media_question']))

        buffer = io.BytesIO()
        document.save(buffer) 
        buffer.seek(0)

        response = StreamingHttpResponse(
            streaming_content=buffer,  
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingm'
        )
        response['Content-Disposition'] = 'attachment;filename=Test.docx'
        response["Content-Encoding"] = 'UTF-8'

        return response


# ------------------------------------
# CATEGORY
# ------------------------------------

class CategoryCreateAPI(APIView):
    
    def post(self, request, format=None):
        serializer = CategorySerializer(data=request.data)

        if serializer.is_valid():
            serializer.save()

            return Response(serializer.data, status=status.HTTP_201_CREATED)

        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CategoryGetAllAPI(APIView):
    
    def get(self, request, *args, **kwargs):
        result = Category.objects.all().order_by('-id').values()

        return Response(result, status=status.HTTP_200_OK)


class CategoryGetAPI(APIView):

    def get(self, request, *args, **kwargs):
        try:
            result = Category.objects.values().get(id=kwargs['pk'])
            return Response(result, status=status.HTTP_200_OK)
        except Category.DoesNotExist:
            raise NotFound(detail=f'Категории с id: {kwargs['pk']} несуществует')


class CategoryGetKeysAPI(APIView):

    def get(self, request, *args, **kwargs):
        ids = kwargs['ids'].split(',')
        result = Category.objects.filter(id__in=ids).values()

        return Response(result, status=status.HTTP_200_OK)


class CategorySearchAPI(ListAPIView):
    serializer_class = CategorySerializer

    def get_queryset(self):
        query = self.request.query_params.get('query', '')

        return Category.objects.filter(
            Q(name__icontains=query),
        )


class CategoryUpdateAPI(UpdateAPIView):

    serializer_class = CategorySerializer
    queryset = Category.objects.all()


class CategoryDeleteAPI(DestroyAPIView):

    def delete(self, request, *args, **kwargs):
        try:
            Category.objects.get(id=kwargs['pk']).delete()
            return HttpResponse(f'Категория с id: {kwargs['pk']} была удалена', status=200)
        except Category.DoesNotExist:
            raise NotFound(detail=f'Категория с id: {kwargs['pk']} не существует.')
